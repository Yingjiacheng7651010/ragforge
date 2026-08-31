"""Unit tests for document parsers and the extension-based registry."""

from pathlib import Path

import pymupdf
import pytest
from docx import Document

from ragforge.core.errors import RAGForgeError
from ragforge.ingestion.parsers import (
    DEFAULT_REGISTRY,
    HtmlParser,
    MarkdownParser,
    ParserRegistry,
    PDFParser,
    WordParser,
)

MARKDOWN_SAMPLE = """# RagForge Sample

## Introduction

This is the intro paragraph.

## Usage

Second paragraph.

| Name | Value |
|------|-------|
| a    | 1     |

```python
# this is code, not a heading
print("hi")
```
"""

HTML_SAMPLE = """<!doctype html>
<html>
<head><title>RagForge Sample</title></head>
<body>
<nav><a href="/x">Nav link</a></nav>
<script>var secret = 1;</script>
<h1>RagForge Sample</h1>
<h2>Introduction</h2>
<p>This is the intro paragraph.</p>
<h2>Usage</h2>
<p>Second paragraph.</p>
<table>
<tr><th>Name</th><th>Value</th></tr>
<tr><td>a</td><td>1</td></tr>
</table>
</body>
</html>
"""


def write_markdown(tmp_path: Path) -> Path:
    path = tmp_path / "sample.md"
    path.write_text(MARKDOWN_SAMPLE, encoding="utf-8")
    return path


def write_html(tmp_path: Path) -> Path:
    path = tmp_path / "sample.html"
    path.write_text(HTML_SAMPLE, encoding="utf-8")
    return path


def write_docx(tmp_path: Path) -> Path:
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("RagForge Sample", level=1)
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph("This is the intro paragraph.")
    doc.add_heading("Usage", level=2)
    doc.add_paragraph("Second paragraph.")
    doc.save(path)
    return path


def write_pdf(tmp_path: Path) -> Path:
    path = tmp_path / "sample.pdf"
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "RagForge Sample", fontsize=24)
    page.insert_text((72, 120), "Introduction", fontsize=16)
    page.insert_text((72, 150), "This is the intro paragraph.", fontsize=12)
    page2 = doc.new_page()
    page2.insert_text((72, 72), "Usage", fontsize=16)
    page2.insert_text((72, 100), "Second paragraph.", fontsize=12)
    doc.save(path)
    doc.close()
    return path


# --- markdown ---


def test_markdown_heading_paths_and_text(tmp_path: Path) -> None:
    document = MarkdownParser().parse(write_markdown(tmp_path))

    assert document.title == "RagForge Sample"
    assert [s.heading_path for s in document.sections] == [
        ["RagForge Sample", "Introduction"],
        ["RagForge Sample", "Usage"],
    ]
    assert document.sections[0].text == "This is the intro paragraph."
    assert document.sections[1].text.startswith("Second paragraph.")
    assert all(section.page is None for section in document.sections)


def test_markdown_table_and_code_fence(tmp_path: Path) -> None:
    document = MarkdownParser().parse(write_markdown(tmp_path))

    assert len(document.tables) == 1
    assert document.tables[0].headers == ["Name", "Value"]
    assert document.tables[0].rows == [["a", "1"]]
    # Code fences join the surrounding section; their contents stay body text.
    assert 'print("hi")' in document.sections[-1].text
    assert not any("# this is code" in " ".join(s.heading_path) for s in document.sections)


# --- pdf ---


def test_pdf_heading_paths_and_pages(tmp_path: Path) -> None:
    document = PDFParser().parse(write_pdf(tmp_path))

    assert document.title == "RagForge Sample"
    assert [s.heading_path for s in document.sections] == [
        ["RagForge Sample", "Introduction"],
        ["RagForge Sample", "Usage"],
    ]
    assert document.sections[0].text == "This is the intro paragraph."
    assert document.sections[0].page == 1
    assert document.sections[1].text == "Second paragraph."
    assert document.sections[1].page == 2
    assert document.metadata["page_count"] == 2


# --- word ---


def test_word_heading_paths_and_tables(tmp_path: Path) -> None:
    document = WordParser().parse(write_docx(tmp_path))

    assert document.title == "RagForge Sample"
    assert [s.heading_path for s in document.sections] == [
        ["RagForge Sample", "Introduction"],
        ["RagForge Sample", "Usage"],
    ]
    assert document.sections[0].text == "This is the intro paragraph."
    assert document.sections[1].text == "Second paragraph."


# --- html ---


def test_html_strips_navigation_and_scripts(tmp_path: Path) -> None:
    document = HtmlParser().parse(write_html(tmp_path))

    all_text = "\n".join(section.text for section in document.sections)
    assert "Nav link" not in all_text
    assert "secret" not in all_text
    assert "This is the intro paragraph." in all_text


def test_html_heading_paths_title_and_tables(tmp_path: Path) -> None:
    document = HtmlParser().parse(write_html(tmp_path))

    assert document.title == "RagForge Sample"
    assert [s.heading_path for s in document.sections] == [
        ["RagForge Sample", "Introduction"],
        ["RagForge Sample", "Usage"],
    ]
    assert len(document.tables) == 1
    assert document.tables[0].headers == ["Name", "Value"]
    assert document.tables[0].rows == [["a", "1"]]


# --- acceptance: same sample across formats ---


@pytest.mark.parametrize("writer", [write_markdown, write_docx, write_pdf])
def test_all_formats_produce_heading_paths(tmp_path: Path, writer: object) -> None:
    path = writer(tmp_path)  # type: ignore[operator]
    document = DEFAULT_REGISTRY.parse(path)

    assert document.title == "RagForge Sample"
    assert document.sections, "expected at least one section"
    paths = [section.heading_path for section in document.sections]
    assert all(len(p) == 2 and p[0] == "RagForge Sample" for p in paths)


# --- registry ---


def test_registry_routes_by_extension() -> None:
    assert isinstance(DEFAULT_REGISTRY.get("file.md"), MarkdownParser)
    assert isinstance(DEFAULT_REGISTRY.get("file.markdown"), MarkdownParser)
    assert isinstance(DEFAULT_REGISTRY.get("file.pdf"), PDFParser)
    assert isinstance(DEFAULT_REGISTRY.get("file.docx"), WordParser)
    assert isinstance(DEFAULT_REGISTRY.get("file.html"), HtmlParser)
    assert isinstance(DEFAULT_REGISTRY.get("file.htm"), HtmlParser)


def test_registry_matches_extension_case_insensitively() -> None:
    assert isinstance(DEFAULT_REGISTRY.get("FILE.PDF"), PDFParser)


def test_registry_unknown_extension_raises() -> None:
    with pytest.raises(RAGForgeError) as exc_info:
        DEFAULT_REGISTRY.get("file.xyz")

    assert exc_info.value.code == "E_UNSUPPORTED_FORMAT"


def test_registry_supports_custom_registration(tmp_path: Path) -> None:
    class TxtParser(MarkdownParser):
        pass

    registry = ParserRegistry()
    registry.register("txt", TxtParser)

    assert isinstance(registry.get("note.txt"), TxtParser)
    with pytest.raises(RAGForgeError):
        registry.get("note.md")  # not registered in this registry
