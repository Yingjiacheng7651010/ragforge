"""Word (.docx) parser based on python-docx: heading styles, paragraphs, tables."""

import re
from pathlib import Path

from docx import Document

from ragforge.ingestion.parsers.base import HeadingStack, ParsedDocument, Parser, Section, Table

#: Matches heading style names like "Heading 1" / "标题 2" (localized Word).
_HEADING_STYLE_RE = re.compile(r"^(?:Heading|标题)\s*(\d+)$", re.IGNORECASE)


class WordParser(Parser):
    """Parse .docx files using paragraph styles to build the heading tree."""

    def parse(self, path: str | Path) -> ParsedDocument:
        source = Path(path)
        document = Document(str(source))

        headings = HeadingStack()
        sections: list[Section] = []
        tables: list[Table] = []
        paragraphs: list[str] = []
        title: str | None = None

        def flush_paragraph() -> None:
            text = "\n\n".join(paragraphs).strip()
            if text:
                sections.append(Section(heading_path=headings.path, text=text))
            paragraphs.clear()

        for paragraph in document.paragraphs:
            style_name = paragraph.style.name if paragraph.style is not None else ""
            match = _HEADING_STYLE_RE.match(style_name or "")
            if match:
                flush_paragraph()
                level = int(match.group(1))
                heading_text = paragraph.text.strip()
                if not heading_text:
                    continue
                headings.push(level, heading_text)
                if title is None and level == 1:
                    title = heading_text
                continue
            if paragraph.text.strip():
                paragraphs.append(paragraph.text.strip())
        flush_paragraph()

        for table in document.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if rows:
                tables.append(Table(headers=list(rows[0]), rows=rows[1:]))

        core_title = document.core_properties.title
        return ParsedDocument(
            doc_id=source.stem,
            title=title or (core_title or "").strip() or source.stem,
            sections=sections,
            tables=tables,
            metadata={"format": "docx", **self._source_metadata(source)},
        )
